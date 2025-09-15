"""
Content extraction implementations for different file types.
"""

try:
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract

    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False

import json
from pathlib import Path
import re
from typing import Any

try:
    import yaml

    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

from .core import Content, ContentExtractor, FileType


class PDFExtractor(ContentExtractor):
    """Extract content from PDF files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.PDF

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract content from PDF file."""
        if not PYPDF_AVAILABLE:
            return Content(
                text="PDF processing not available",
                metadata={"error": "pypdf not installed"},
                file_type=FileType.PDF,
                confidence=0.0,
                extraction_method="pdf_extractor",
            )
        try:
            with open(file_path, "rb") as file:
                reader = pypdf.PdfReader(file)

                # Get first page content
                first_page = reader.pages[0].extract_text()

                # Get metadata
                metadata: dict[str, Any] = reader.metadata or {}

                # Extract key information
                title = metadata.get("/Title", "")
                author = metadata.get("/Author", "")
                subject = metadata.get("/Subject", "")
                page_count = len(reader.pages)

                # Limit content based on config
                max_chars = config.get("max_chars", 1000)
                content_text = first_page[:max_chars]

                return Content(
                    text=content_text,
                    metadata={
                        "title": title,
                        "author": author,
                        "subject": subject,
                        "page_count": page_count,
                        "file_type": "pdf",
                    },
                    file_type=FileType.PDF,
                    confidence=0.9,
                    extraction_method="pdf_extractor",
                )
        except Exception as e:
            return Content(
                text=f"Error reading PDF: {str(e)}",
                metadata={"error": str(e)},
                file_type=FileType.PDF,
                confidence=0.1,
                extraction_method="pdf_extractor",
            )


class DOCXExtractor(ContentExtractor):
    """Extract content from DOCX files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.DOCX

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract content from DOCX file."""
        if not DOCX_AVAILABLE:
            return Content(
                text="DOCX processing not available",
                metadata={"error": "python-docx not installed"},
                file_type=FileType.DOCX,
                confidence=0.0,
                extraction_method="docx_extractor",
            )
        try:
            doc = Document(str(file_path))

            # Extract headings
            headings = [
                p.text
                for p in doc.paragraphs
                if p.style and p.style.name and p.style.name.startswith("Heading")
            ]

            # Extract first few paragraphs
            content_paragraphs = [p.text for p in doc.paragraphs[:5]]
            content_text = "\n".join(content_paragraphs)

            # Extract metadata
            metadata = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "keywords": doc.core_properties.keywords,
                "file_type": "docx",
            }

            # Limit content based on config
            max_chars = config.get("max_chars", 1000)
            content_text = content_text[:max_chars]

            return Content(
                text=content_text,
                metadata={**metadata, "headings": headings[:3]},  # Top 3 headings
                file_type=FileType.DOCX,
                confidence=0.9,
                extraction_method="docx_extractor",
            )
        except Exception as e:
            return Content(
                text=f"Error reading DOCX: {str(e)}",
                metadata={"error": str(e)},
                file_type=FileType.DOCX,
                confidence=0.1,
                extraction_method="docx_extractor",
            )


class CodeExtractor(ContentExtractor):
    """Extract content from code files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.CODE

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract content from code file."""
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            # Extract key elements
            functions = self._find_functions(content)
            classes = self._find_classes(content)
            imports = self._find_imports(content)
            comments = self._find_comments(content)

            # Determine language
            language = self._guess_language(file_path, content)

            # Create summary content
            summary_parts = []
            if functions:
                summary_parts.append(f"Functions: {', '.join(functions[:3])}")
            if classes:
                summary_parts.append(f"Classes: {', '.join(classes[:3])}")
            if imports:
                summary_parts.append(f"Imports: {', '.join(imports[:3])}")
            if comments:
                summary_parts.append(f"Comments: {comments[0][:100]}...")

            # Limit content based on config
            max_chars = config.get("max_chars", 1000)
            content_text = content[:max_chars]

            return Content(
                text=content_text,
                metadata={
                    "functions": functions[:5],
                    "classes": classes[:5],
                    "imports": imports[:5],
                    "comments": comments[:2],
                    "language": language,
                    "file_type": "code",
                },
                file_type=FileType.CODE,
                confidence=0.8,
                extraction_method="code_extractor",
            )
        except Exception as e:
            return Content(
                text=f"Error reading code file: {str(e)}",
                metadata={"error": str(e)},
                file_type=FileType.CODE,
                confidence=0.1,
                extraction_method="code_extractor",
            )

    def _find_functions(self, content: str) -> list[str]:
        """Find function definitions."""
        patterns = [
            r"def\s+(\w+)",  # Python
            r"function\s+(\w+)",  # JavaScript
            r"func\s+(\w+)",  # Go, Swift
            r"fn\s+(\w+)",  # Rust
            r"(\w+)\s*\([^)]*\)\s*{",  # Generic function calls
        ]

        functions = []
        for pattern in patterns:
            matches = re.findall(pattern, content)
            functions.extend(matches)

        return list(set(functions))  # Remove duplicates

    def _find_classes(self, content: str) -> list[str]:
        """Find class definitions."""
        patterns = [
            r"class\s+(\w+)",  # Most languages
            r"struct\s+(\w+)",  # C/C++, Rust, Go
            r"interface\s+(\w+)",  # Java, TypeScript
            r"trait\s+(\w+)",  # Rust, Scala
        ]

        classes = []
        for pattern in patterns:
            matches = re.findall(pattern, content)
            classes.extend(matches)

        return list(set(classes))

    def _find_imports(self, content: str) -> list[str]:
        """Find import statements."""
        patterns = [
            r"import\s+([^\n;]+)",  # Python, Java, etc.
            r"from\s+(\w+)",  # Python
            r'require\s*\(["\']([^"\']+)["\']',  # Node.js
            r"use\s+([^\n;]+)",  # Rust, PHP
            r'#include\s*[<"]([^>"]+)[>"]',  # C/C++
        ]

        imports = []
        for pattern in patterns:
            matches = re.findall(pattern, content)
            imports.extend(matches)

        return list(set(imports))

    def _find_comments(self, content: str) -> list[str]:
        """Find comment blocks."""
        patterns = [
            r"/\*.*?\*/",  # Block comments
            r"//.*",  # Line comments
            r"#.*",  # Hash comments
        ]

        comments = []
        for pattern in patterns:
            matches = re.findall(pattern, content, re.DOTALL)
            comments.extend(matches)

        return comments

    def _guess_language(self, file_path: Path, content: str) -> str:
        """Guess programming language."""
        extension = file_path.suffix.lower()

        lang_map = {
            ".py": "Python",
            ".js": "JavaScript",
            ".ts": "TypeScript",
            ".java": "Java",
            ".cpp": "C++",
            ".c": "C",
            ".go": "Go",
            ".rs": "Rust",
            ".rb": "Ruby",
            ".php": "PHP",
            ".swift": "Swift",
            ".kt": "Kotlin",
            ".scala": "Scala",
            ".hs": "Haskell",
            ".ml": "OCaml",
            ".fs": "F#",
            ".clj": "Clojure",
            ".r": "R",
        }

        if extension in lang_map:
            return lang_map[extension]

        # Content-based detection
        if "def " in content and "import " in content:
            return "Python"
        elif "function" in content and "var " in content:
            return "JavaScript"
        elif "func " in content and "package " in content:
            return "Go"
        elif "fn " in content and "use " in content:
            return "Rust"

        return "Unknown"


class TextExtractor(ContentExtractor):
    """Extract content from text files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type in [
            FileType.TEXT,
            FileType.MARKUP,
            FileType.CONFIG,
            FileType.DATA,
        ]

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract content from text file."""
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            # Determine specific type
            specific_type = self._determine_text_type(content, file_path)

            # Extract relevant information based on type
            metadata = self._extract_metadata(content, specific_type)

            # Limit content based on config
            max_chars = config.get("max_chars", 1000)
            content_text = content[:max_chars]

            return Content(
                text=content_text,
                metadata=metadata,
                file_type=FileType.TEXT,
                confidence=0.8,
                extraction_method="text_extractor",
            )
        except Exception as e:
            return Content(
                text=f"Error reading text file: {str(e)}",
                metadata={"error": str(e)},
                file_type=FileType.TEXT,
                confidence=0.1,
                extraction_method="text_extractor",
            )

    def _determine_text_type(self, content: str, file_path: Path) -> str:
        """Determine specific text file type."""
        if content.strip().startswith(("{", "[")):
            return "json"
        elif content.strip().startswith("---") or ":" in content and "\n" in content:
            return "yaml"
        elif "<" in content and ">" in content:
            return "markup"
        elif "=" in content and "\n" in content:
            return "config"
        else:
            return "text"

    def _extract_metadata(self, content: str, text_type: str) -> dict[str, Any]:
        """Extract metadata based on text type."""
        metadata: dict[str, Any] = {"file_type": text_type}

        if text_type == "json":
            try:
                data = json.loads(content)
                if isinstance(data, dict):
                    metadata["keys"] = list(data.keys())[:5]
            except Exception:
                pass
        elif text_type == "yaml" and YAML_AVAILABLE:
            try:
                data = yaml.safe_load(content)
                if isinstance(data, dict):
                    metadata["keys"] = list(data.keys())[:5]
            except Exception:
                pass
        elif text_type == "markup":
            # Extract headings
            headings = re.findall(
                r"<h[1-6][^>]*>(.*?)</h[1-6]>", content, re.IGNORECASE
            )
            metadata["headings"] = headings[:3]

        return metadata


class ImageExtractor(ContentExtractor):
    """Extract content from image files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.IMAGE

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract content from image file."""
        if not IMAGE_AVAILABLE:
            return Content(
                text="Image processing not available",
                metadata={"error": "Pillow/pytesseract not installed"},
                file_type=FileType.IMAGE,
                confidence=0.0,
                extraction_method="image_extractor",
            )
        try:
            # Get image metadata
            with Image.open(file_path) as img:
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode,
                    "file_type": "image",
                }

                # Try OCR if enabled
                ocr_text = ""
                if config.get("enable_ocr", True):
                    try:
                        ocr_text = pytesseract.image_to_string(img)
                    except Exception:
                        pass

            return Content(
                text=ocr_text[:500],  # Limit OCR text
                metadata=metadata,
                file_type=FileType.IMAGE,
                confidence=0.7,
                extraction_method="image_extractor",
            )
        except Exception as e:
            return Content(
                text=f"Error reading image: {str(e)}",
                metadata={"error": str(e)},
                file_type=FileType.IMAGE,
                confidence=0.1,
                extraction_method="image_extractor",
            )


class FallbackExtractor(ContentExtractor):
    """Fallback extractor for unknown file types."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.UNKNOWN

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Extract basic information from unknown file."""
        try:
            stat = file_path.stat()
            return Content(
                text=f"File: {file_path.name}",
                metadata={
                    "size": stat.st_size,
                    "modified": stat.st_mtime,
                    "file_type": "unknown",
                },
                file_type=FileType.UNKNOWN,
                confidence=0.3,
                extraction_method="fallback_extractor",
            )
        except Exception as e:
            return Content(
                text=f"File: {file_path.name}",
                metadata={"error": str(e)},
                file_type=FileType.UNKNOWN,
                confidence=0.1,
                extraction_method="fallback_extractor",
            )


class CompositeExtractor(ContentExtractor):
    """Composite extractor that tries multiple extraction methods."""

    def __init__(self, extractors: list[ContentExtractor]):
        self.extractors = extractors

    def can_handle(self, file_type: FileType) -> bool:
        return any(extractor.can_handle(file_type) for extractor in self.extractors)

    def extract(self, file_path: Path, config: dict[str, Any]) -> Content:
        """Try all extractors that can handle the file type."""
        for extractor in self.extractors:
            if extractor.can_handle(config.get("file_type", FileType.UNKNOWN)):
                try:
                    return extractor.extract(file_path, config)
                except Exception:
                    continue

        # If all extractors fail, use fallback
        return FallbackExtractor().extract(file_path, config)
